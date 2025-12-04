import pandas as pd
import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Mm, RGBColor
from docx.enum.style import WD_STYLE_TYPE
import fitz


class ReportCreator:
    def __init__(self, name_file, path_to_data=''):
        self.name_file = name_file
        self.df = pd.read_csv(f"{path_to_data}{name_file}.csv")
        self.total_amount = (self.df['Количество'] * self.df['Цена']).sum()

    @staticmethod
    def match_width_of_content(name_file, output_path=''):
        full_path = f"{output_path}{name_file}.xlsx"
        wb = openpyxl.load_workbook(full_path)
        ws = wb.active
        for col in ws.columns:
            max_length = 0
            column_name = col[0].column_letter
            for cell in col:
                if cell.value:
                    length = len(str(cell.value))
                    if length > max_length:
                        max_length = length
            ws.column_dimensions[column_name].width = max_length + 2
        wb.save(full_path)

    def create_report_excel(self, color, value_for_condition, column_names, name_file=None, output_path=''):
        if column_names is not list:
            column_names = [column_names]
        if name_file is None:
            name_file = self.name_file

        formatted_data = self.df.style.apply(lambda x: [f"background-color: {color}" if value > value_for_condition else "" for value in x],
                                        subset=column_names)
        formatted_data.to_excel(f"{output_path}{name_file}.xlsx", index=False, engine='openpyxl')
        self.match_width_of_content(name_file, output_path)

    def create_report_word(self, name_file=None, output_path=''):
        if name_file is None:
            name_file = self.name_file

        doc = Document()
        style_font = doc.styles['Normal'].font
        style_font.name = "Times New Roman"
        style_font.size = Pt(12)
        style_my_heading = doc.styles.add_style("My Heading", WD_STYLE_TYPE.PARAGRAPH)
        style_my_heading_font = style_my_heading.font
        style_my_heading_font.name = "Times New Roman"
        style_my_heading_font.color.rgb = RGBColor(0, 0, 0)
        style_my_heading_font.size = Pt(14)
        style_my_heading_font.bold = True
        style_my_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("Отчёт по продажам", style="My Heading")
        doc.add_paragraph("Основные позиции:")
        for el in self.df['Имя']:
            doc.add_paragraph(el, style="List Bullet")
        doc.add_paragraph(f"Итоговая сумма – {self.total_amount} рублей.")
        doc.save(f"{output_path}{name_file}.docx")

    def create_report_pdf(self, date, originator, fontname="Times-Bold", name_file=None, output_path=''):
        if name_file is None:
            name_file = self.name_file

        pdf = fitz.open("Template.pdf")
        page = pdf[0]

        page.insert_text((125, 110), str(date), fontsize=12, fontname=fontname)
        page.insert_text((445, 110), originator, fontsize=12, fontname=fontname)

        start_y = 152
        height_cell = 18.7197265625
        for idx in range(len(self.df)):
            page.insert_text((210, start_y), self.df.iat[idx, 0], fontsize=8, fontname=fontname)
            page.insert_text((320, start_y), str(self.df.iat[idx, 1]), fontsize=10, fontname=fontname)
            page.insert_text((445, start_y), str(self.df.iat[idx, 2]), fontsize=10, fontname=fontname)
            start_y += height_cell
        page.insert_text((453, 724), str(self.total_amount), fontsize=12, fontname=fontname)
        pdf.save(f"{output_path}{name_file}.pdf")
        pdf.close()


if __name__ == "__main__":
    name = "Sale of equipment"
    report_creator = ReportCreator(name)
    report_creator.create_report_excel("green", 100, "Количество")
    report_creator.create_report_word()
    report_creator.create_report_pdf("20.11.2025", "Nikolay")